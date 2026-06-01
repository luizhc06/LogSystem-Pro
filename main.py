"""
LogSystem Pro — backend principal.

Responsabilidades:
  - Abrir a janela PyWebView com app.html
  - Expor a classe API pro JavaScript (login, pedidos, import/export, backup)
  - Traduzir status entre planilha Excel e o formato interno do sistema
  - Rodar backup automático em thread separada

Fluxo: LogSystem.bat/sh → main.py → webview.start() → app.html chama api.*
"""

import os
import re
import sys
import json
import csv
import io
import threading
import time
import glob
import hashlib
import secrets
from datetime import datetime, date, timedelta

from modules import auth, bootstrap, database

APP_VERSION = "28"

# Caminhos — funciona tanto rodando .py quanto empacotado em .exe (PyInstaller)
def _app_base_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def _resource_path(name):
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        bundled = os.path.join(sys._MEIPASS, name)
        if os.path.exists(bundled):
            return bundled
    return os.path.join(_app_base_dir(), name)

BASE_DIR = _app_base_dir()
HTML_FILE = _resource_path("app.html")
DATA_DIR    = os.path.join(BASE_DIR, "data")
DATA_FILE   = os.path.join(DATA_DIR, "data.json")
CONFIG_FILE = os.path.join(BASE_DIR, "config.json")
BACKUP_DIR  = os.path.join(DATA_DIR, "backups")
USERS_FILE  = os.path.join(DATA_DIR, "users.json")
SESSION_FILE = os.path.join(DATA_DIR, ".session")

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(BACKUP_DIR, exist_ok=True)

# Tradução de rótulos da planilha Excel ↔ IDs usados no app (app.html → STATUS)
STATUS_XLSX_TO_SYS = {
    "NO PRAZO": "no_prazo", "ATRASADO": "em_atraso", "ATRASO": "em_atraso",
    "RISCO DE ATRASO": "risco_atraso", "TRANSPOR. ATRASADA": "em_atraso",
    "EM TRANSITO": "em_transito", "EM TRÂNSITO": "em_transito",
    "REENVIO": "reenvio",
    "SAIU PARA ENTREGA": "em_rota", "SAIU PARA ENTREGA": "em_rota",
    "ENTREGUE NO PRAZO": "entregue_prazo", "ENTREGUE FORA DO PRAZO": "entregue_atraso",
    "ENTREGUE": "entregue_prazo",
    "NAO ENTREGUE": "nao_entregue", "NÃO ENTREGUE": "nao_entregue",
    "EXTRAVIADO": "extraviado", "EXTRAVIO": "extraviado",
    "RETORNO": "retorno", "DEVOLUCAO": "retorno", "DEVOLUÇÃO": "retorno",
    "RETORNO ENTREGUE": "retorno_entregue", "RETORNO CONCLUIDO": "retorno_entregue",
    "REENVIO ENTREGUE": "reenvio_prazo", "REENVIO FORA DO PRAZO": "reenvio_atraso",
}
STATUS_FINAL_TO_SYS = {
    "EM VIAGEM": "em_transito",
    "ENTREGUE": "entregue_prazo",
    "ENTREGUE NO PRAZO": "entregue_prazo",
    "ENTREGUE FORA DO PRAZO": "entregue_atraso",
    "NAO ENTREGUE": "nao_entregue", "NÃO ENTREGUE": "nao_entregue",
    "NAO ENTREGUE": "nao_entregue",
}
STATUS_SYS_TO_XLSX = {
    "no_prazo":         ("NO PRAZO",                "EM VIAGEM"),
    "em_atraso":        ("ATRASADO",                "EM VIAGEM"),
    "risco_atraso":     ("RISCO DE ATRASO",         "EM VIAGEM"),
    "em_transito":      ("NO PRAZO",                "EM VIAGEM"),
    "em_rota":          ("SAIU PARA ENTREGA",      "EM VIAGEM"),
    "reenvio":          ("REENVIO",                 "EM VIAGEM"),
    "nao_entregue":     ("ATRASADO",                "NÃO ENTREGUE"),
    "extraviado":       ("EXTRAVIADO",              "EXTRAVIADO"),
    "retorno":          ("RETORNO",                 "RETORNO"),
    "retorno_entregue": ("RETORNO ENTREGUE",         "RETORNO ENTREGUE"),
    "entregue_prazo":   ("ENTREGUE NO PRAZO",       "ENTREGUE"),
    "entregue_atraso":  ("ENTREGUE FORA DO PRAZO",  "ENTREGUE"),
    "reenvio_prazo":    ("REENVIO ENTREGUE",        "ENTREGUE"),
    "reenvio_atraso":   ("REENVIO FORA DO PRAZO",   "ENTREGUE"),
}

# Login — delega pro modules/auth.py; sessão fica só em memória até fechar o app
def hash_password(password: str) -> str:
    return auth.hash_password(password)

def load_users() -> dict:
    return auth.load_users(USERS_FILE)

def save_users(users: dict):
    auth.save_users(USERS_FILE, users)

_active_session = {"token": None, "user": None, "email": None}  # zera ao reiniciar

def generate_token() -> str:
    return secrets.token_hex(32)

# Abrir/salvar arquivo — tenta tkinter, cai pro dialog nativo do pywebview
def open_file_dialog(title="Selecionar arquivo", filetypes=None, save=False, save_name="arquivo.xlsx"):
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk(); root.withdraw()
        root.attributes('-topmost', True); root.update()
        if save:
            path = filedialog.asksaveasfilename(
                title=title, initialfile=save_name,
                defaultextension=os.path.splitext(save_name)[1],
                filetypes=filetypes or [("Todos","*.*")], parent=root)
        else:
            path = filedialog.askopenfilename(
                title=title, filetypes=filetypes or [("Todos","*.*")], parent=root)
        root.destroy()
        return path if path else None
    except:
        try:
            if save:
                r = webview.windows[0].create_file_dialog(webview.SAVE_DIALOG,
                    directory=os.path.expanduser("~"), save_filename=save_name)
            else:
                r = webview.windows[0].create_file_dialog(webview.OPEN_DIALOG)
            if r: return r[0] if isinstance(r,(list,tuple)) else r
        except: pass
        return None

# Conversões seguras — planilha manda número serial, string vazia, ISO, etc.
def excel_serial_to_date(val):
    if isinstance(val, (datetime, date)):
        return val if isinstance(val, datetime) else datetime(val.year, val.month, val.day)
    if isinstance(val, (int, float)) and 30000 < val < 60000:
        try: return datetime(1899, 12, 30) + timedelta(days=int(val))
        except: return None
    if isinstance(val, str):
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
            try: return datetime.strptime(val.strip(), fmt)
            except: pass
    return None

def safe_float(v):
    """Convert any value to float safely — handles comma decimal separator"""
    if v is None or str(v).strip() == '': return 0.0
    try: return float(str(v).replace(',', '.').strip())
    except: return 0.0

def safe_str(val, default=""):
    """Converte célula Excel para string.
    float 152247010410.0 → '152247010410' (sem .0 que corromperia IDs Magalu)."""
    if val is None: return default
    if isinstance(val, (datetime, date)): return val.strftime("%d/%m/%Y")
    if isinstance(val, float) and val == int(val):
        return str(int(val)).strip()
    return str(val).strip()

def safe_date_iso(val):
    d = excel_serial_to_date(val)
    return d.strftime("%Y-%m-%dT12:00:00") if d else ""


# Pedidos e config — leitura/gravação via modules/database.py (SQLite)
def load_data():
    """Carrega todos os pedidos do SQLite (com fallback para data.json)."""
    try:
        return database.listar_pedidos()
    except Exception as exc:
        print(f"[ERRO] Falha ao carregar pedidos do banco: {exc}")
        if os.path.exists(DATA_FILE):
            try:
                with open(DATA_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, list):
                    return data
            except Exception as file_exc:
                print(f"[ERRO] Fallback data.json também falhou: {file_exc}")
        return []

def save_data(orders):
    """Salva todos os pedidos no SQLite e sincroniza data.json."""
    try:
        database.salvar_pedidos(orders)
    except Exception as exc:
        print(f"[ERRO] Falha ao salvar no banco: {exc}")
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(orders, f, ensure_ascii=False, indent=2)

def load_config():
    """Carrega transportadoras/plataformas customizadas."""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except: pass
    return {}

def save_config(cfg):
    """Salva transportadoras/plataformas no arquivo de config."""
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Config save error: {e}")

# Snapshot JSON em data/backups/ — limpa arquivos com mais de 24h
def do_backup():
    try:
        caminho = database.backup_para_json()
        # Limpa backups antigos (mantém últimos 24h)
        cutoff = time.time() - 86400
        for bf in glob.glob(os.path.join(BACKUP_DIR, "backup_*.json")):
            if os.path.getmtime(bf) < cutoff:
                try: os.remove(bf)
                except: pass
    except Exception as e:
        print(f"Backup error: {e}")

def backup_loop():
    while True:
        time.sleep(300)
        do_backup()

_idc = 10000
def next_id():
    global _idc; _idc += 1; return f"ORD-{_idc:05d}"

# Achar pedido duplicado na importação (mesma NF, chave NF-e ou nº pedido)
def normalize_nf(nf: str) -> str:
    """Normalize NF for comparison: remove NF-/NF prefix, strip, lowercase."""
    s = str(nf or "").strip()
    s = re.sub(r"^NF-?", "", s, flags=re.IGNORECASE).strip()
    return s.lower()

def find_existing_order(orders, nf=None, chave_nfe=None, num_pedido=None):
    """Find existing order by chave NF-e (primary), then NF number."""
    if chave_nfe:
        for o in orders:
            if o.get("chavNfe") and o["chavNfe"].strip() == chave_nfe.strip():
                return o
    if nf:
        n1 = normalize_nf(nf)
        for o in orders:
            if normalize_nf(o.get("nf","")) == n1:
                return o
    return None

def _merge_backup_raw(raw):
    """Mescla backup JSON com dados atuais. Retorna dict resultado."""
    if isinstance(raw, list):
        backup_data = raw
        backup_cfg = {}
    elif isinstance(raw, dict):
        backup_data = raw.get("orders", [])
        backup_cfg = raw.get("config", {})
    else:
        raise ValueError("Formato de backup invalido")

    current_data = load_data()

    config_restored = False
    if backup_cfg:
        current_cfg = load_config()
        merged_cfg = dict(current_cfg)
        for key in ("transportadoras", "plataformas"):
            bk_list = backup_cfg.get(key, [])
            cur_list = current_cfg.get(key, [])
            if bk_list:
                merged = list(cur_list)
                for item in bk_list:
                    if item not in merged:
                        merged.append(item)
                merged_cfg[key] = merged
                config_restored = True
        if config_restored:
            save_config(merged_cfg)

    current_by_id = {str(o.get("id", "")): o for o in current_data if o.get("id")}
    current_by_nf = {
        str(o.get("nf", "")).replace("NF-", "").strip(): o
        for o in current_data if o.get("nf")
    }

    added, merged_chamados, merged_hist = 0, 0, 0
    for bk in backup_data:
        oid = str(bk.get("id", ""))
        nf = str(bk.get("nf", "")).replace("NF-", "").strip()

        cur = None
        if oid and oid in current_by_id:
            cur = current_by_id[oid]
        elif nf and nf in current_by_nf:
            cur = current_by_nf[nf]

        if cur is None:
            current_data.append(bk)
            added += 1
            continue

        cur_chamado_ids = {str(ch.get("id", "")) for ch in (cur.get("chamados") or [])}
        for ch in (bk.get("chamados") or []):
            cid = str(ch.get("id", ""))
            if cid and cid not in cur_chamado_ids:
                if not isinstance(cur.get("chamados"), list):
                    cur["chamados"] = []
                cur["chamados"].append(ch)
                cur_chamado_ids.add(cid)
                merged_chamados += 1

        cur_hist_keys = {
            f"{h.get('data', '')}|{h.get('acao', '')}"
            for h in (cur.get("historico") or [])
        }
        for h in (bk.get("historico") or []):
            key = f"{h.get('data', '')}|{h.get('acao', '')}"
            if key not in cur_hist_keys:
                if not isinstance(cur.get("historico"), list):
                    cur["historico"] = []
                cur["historico"].append(h)
                cur_hist_keys.add(key)
                merged_hist += 1

        CAMPOS_COMPLEMENTAVEIS = [
            "chavNfe", "rastreio", "cidade", "uf", "numeroPedido",
            "dataEntregaReal", "diasEntregaReal", "clienteAvisado",
            "dataContato", "responsavel", "nfOriginal", "nfReenvioVinculada",
            "custoFrete", "custoReenvio", "custoExtra", "valor", "peso", "volumes",
        ]
        for k in CAMPOS_COMPLEMENTAVEIS:
            if not cur.get(k) and bk.get(k):
                cur[k] = bk[k]

    save_data(current_data)
    parts = [f"{added} pedido(s) restaurado(s)"]
    if merged_chamados:
        parts.append(f"{merged_chamados} chamado(s) recuperado(s)")
    if merged_hist:
        parts.append(f"{merged_hist} entrada(s) de historico recuperada(s)")
    return {
        "ok": True,
        "total": len(current_data),
        "added": added,
        "merged_chamados": merged_chamados,
        "merged_hist": merged_hist,
        "msg": " | ".join(parts),
    }


# Ponte Python ↔ JavaScript — cada método vira window.pywebview.api.nome()
class API:
    # login, sessão, troca de senha
    def login(self, email: str, password: str) -> dict:
        users = load_users()
        email_key = email.strip().lower()
        u = users.get(email_key)
        if not u:
            return {"ok": False, "msg": "Email nao encontrado"}
        if not auth.verify_password(password, u.get("password_hash", "")):
            return {"ok": False, "msg": "Senha incorreta"}
        token = generate_token()
        _active_session["token"] = token
        _active_session["user"] = u["name"]
        _active_session["email"] = email
        return {"ok": True, "token": token, "name": u["name"], "email": email, "role": u.get("role","user")}

    def verify_session(self, token: str) -> dict:
        if _active_session["token"] and _active_session["token"] == token:
            return {"ok": True, "name": _active_session["user"], "email": _active_session["email"]}
        return {"ok": False}

    def logout(self, token: str) -> dict:
        if _active_session["token"] == token:
            _active_session["token"] = None
            _active_session["user"] = None
        return {"ok": True}

    def change_password(self, token: str, old_pw: str, new_pw: str) -> dict:
        if not _active_session["token"] or _active_session["token"] != token:
            return {"ok": False, "msg": "Nao autenticado"}
        users = load_users()
        email = _active_session["email"]
        u = users.get(email)
        if not u or not auth.verify_password(old_pw, u.get("password_hash", "")):
            return {"ok": False, "msg": "Senha atual incorreta"}
        if len(new_pw) < 4:
            return {"ok": False, "msg": "Nova senha muito curta (minimo 4 caracteres)"}
        u["password_hash"] = hash_password(new_pw)
        save_users(users)
        return {"ok": True}

    def _check_auth(self, token: str) -> bool:
        return bool(_active_session["token"] and _active_session["token"] == token)

    # leitura/gravação de pedidos
    def get_orders(self):
        return load_data()

    def save_orders(self, orders):
        save_data(orders)
        return {"ok": True}

    # listar, criar e restaurar backups
    def manual_backup(self):
        do_backup()
        bfiles = sorted(glob.glob(os.path.join(BACKUP_DIR, "backup_*.json")))
        return {"ok": True, "count": len(bfiles),
                "latest": os.path.basename(bfiles[-1]) if bfiles else ""}

    def list_backups(self):
        bfiles = sorted(glob.glob(os.path.join(BACKUP_DIR, "backup_*.json")), reverse=True)
        return [{"name": os.path.basename(f),
                 "size": os.path.getsize(f),
                 "date": datetime.fromtimestamp(os.path.getmtime(f)).strftime("%d/%m/%Y %H:%M:%S")}
                for f in bfiles[:30]]

    def restore_backup(self, name):
        """Restaura backup incremental a partir da pasta backups/."""
        try:
            path = os.path.join(BACKUP_DIR, name)
            if not os.path.exists(path):
                return {"ok": False, "msg": "Backup nao encontrado"}
            with open(path, "r", encoding="utf-8") as f:
                raw = json.load(f)
            return _merge_backup_raw(raw)
        except Exception as e:
            return {"ok": False, "msg": str(e)}

    def restore_backup_content(self, content: str):
        """Restaura backup a partir do conteudo JSON (arrastar arquivo)."""
        try:
            raw = json.loads(content)
            return _merge_backup_raw(raw)
        except json.JSONDecodeError:
            return {"ok": False, "msg": "Arquivo JSON invalido"}
        except Exception as e:
            return {"ok": False, "msg": str(e)}

    def get_config(self):
        """Retorna configurações salvas (transportadoras, plataformas)."""
        return load_config()

    def save_config_data(self, cfg):
        """Salva configurações (transportadoras, plataformas) em config.json."""
        try:
            save_config(cfg)
            return {"ok": True}
        except Exception as e:
            return {"ok": False, "msg": str(e)}

    # relatórios Word e Excel de chamados
    def generate_report(self, orders):
        """Gera relatório DOCX de Chamados e Observações de todos os pedidos."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime as dt

            doc = Document()

            # Estilos
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(10)

            # Título
            title = doc.add_heading('Relatório de Chamados e Observações', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub = doc.add_paragraph(f'LogSystem Pro — Gerado em {dt.now().strftime("%d/%m/%Y %H:%M")}')
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.size = Pt(9)
            sub.runs[0].font.color.rgb = RGBColor(0x78,0x91,0xAA)
            doc.add_paragraph('')

            # Filtrar pedidos com chamados ou obs
            with_data = [o for o in (orders or [])
                         if (o.get('chamados') and len(o['chamados'])>0)
                         or (o.get('observacoes') and len(o['observacoes'])>0)]

            if not with_data:
                doc.add_paragraph('Nenhum pedido com chamados ou observações registradas.')
            else:
                for o in with_data:
                    nf   = str(o.get('nf',''))
                    cli  = str(o.get('cliente',''))
                    sts  = str(o.get('status',''))
                    trans = str(o.get('transportadora',''))
                    plat  = str(o.get('plataforma',''))

                    # Cabeçalho do pedido
                    h = doc.add_heading(f'{nf} — {cli}', level=1)
                    h.runs[0].font.size = Pt(13)
                    info = doc.add_paragraph()
                    info.add_run(f'Status: {sts}   |   {trans}   |   {plat}').font.size = Pt(9)

                    # ── CHAMADOS
                    chamados = o.get('chamados') or []
                    if chamados:
                        doc.add_heading('Chamados', level=2)
                        for ch in chamados:
                            p = doc.add_paragraph(style='List Bullet')
                            run = p.add_run(f'[{ch.get("status","").upper()}] {ch.get("tipo","")} — {ch.get("data","")}')
                            run.bold = True
                            run.font.size = Pt(10)
                            if ch.get('descricao'):
                                d2 = doc.add_paragraph(ch['descricao'])
                                d2.paragraph_format.left_indent = Inches(0.3)
                                d2.runs[0].font.size = Pt(9)
                            if ch.get('resolucao'):
                                r2 = doc.add_paragraph(f'Resolução: {ch["resolucao"]}')
                                r2.paragraph_format.left_indent = Inches(0.3)
                                r2.runs[0].font.size = Pt(9)
                                r2.runs[0].font.italic = True

                    # ── OBSERVAÇÕES
                    obs_list = o.get('observacoes') or []
                    if obs_list:
                        doc.add_heading('Observações', level=2)
                        for ob in obs_list:
                            p = doc.add_paragraph(style='List Bullet')
                            run = p.add_run(f'{ob.get("data","")} — {ob.get("usuario","Operador")}')
                            run.bold = True
                            run.font.size = Pt(9)
                            if ob.get('texto'):
                                d2 = doc.add_paragraph(ob['texto'])
                                d2.paragraph_format.left_indent = Inches(0.3)
                                d2.runs[0].font.size = Pt(10)

                    doc.add_paragraph('─' * 60).runs[0].font.color.rgb = RGBColor(0x2a,0x3a,0x55)

            # Sumário no início
            ts = dt.now().strftime('%Y%m%d_%H%M%S')
            out_path = os.path.join(BASE_DIR, f'relatorio_chamados_{ts}.docx')
            doc.save(out_path)
            return {'ok': True, 'path': out_path, 'total': len(with_data)}
        except Exception as e:
            import traceback
            return {'ok': False, 'msg': str(e), 'trace': traceback.format_exc()}


    def export_chamados_report(self):
        """Exporta relatorio completo de chamados em XLSX."""
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            return {"ok": False, "msg": "openpyxl nao instalado"}
        try:
            fname = f"chamados_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            path = open_file_dialog("Salvar relatorio de chamados",
                filetypes=[("Excel","*.xlsx")], save=True, save_name=fname)
            if not path: return {"ok": False, "msg": "Cancelado"}

            orders = load_data()
            wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Chamados"
            hfill = PatternFill("solid", fgColor="1F3864")
            hfont = Font(bold=True, color="FFFFFF", size=10)
            thin = Side(border_style="thin", color="CCCCCC")
            bdr = Border(left=thin, right=thin, top=thin, bottom=thin)
            HEADERS = ["NF","CLIENTE","TITULO CHAMADO","STATUS","DATA ABERTURA","DATA CONCLUSAO",
                       "RESPONSAVEL","DESCRICAO","HISTORICO RESUMIDO","TAREFAS","COMENTARIOS"]
            for c, h in enumerate(HEADERS, 1):
                cell = ws.cell(row=1, column=c, value=h)
                cell.fill=hfill; cell.font=hfont
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border=bdr
            ws.row_dimensions[1].height = 28

            STATUS_FILLS = {
                "aberto": PatternFill("solid",fgColor="FFC7CE"),
                "em andamento": PatternFill("solid",fgColor="FFEB9C"),
                "resolvido": PatternFill("solid",fgColor="C6EFCE"),
            }
            row_idx = 2
            for o in orders:
                for ch in (o.get("chamados") or []):
                    hist_summary = " | ".join(
                        h.get("acao","") for h in (ch.get("historico") or [])[-5:])
                    tasks = "; ".join(
                        f"[{'x' if t.get('done') else ' '}] {t.get('texto','')}"
                        for t in (ch.get("tarefas") or []))
                    comments = " | ".join(
                        f"{c.get('usuario','')}: {c.get('texto','')}"
                        for c in (ch.get("comentarios") or []))
                    row_data = [
                        o.get("nf",""), o.get("cliente",""),
                        ch.get("titulo",""), ch.get("status",""),
                        ch.get("data",""), ch.get("dataConclusao",""),
                        ch.get("responsavel",""), ch.get("descricao",""),
                        hist_summary, tasks, comments,
                    ]
                    sfill = STATUS_FILLS.get(ch.get("status","aberto"))
                    for ci, val in enumerate(row_data, 1):
                        cell = ws.cell(row=row_idx, column=ci, value=val)
                        cell.border = bdr
                        cell.alignment = Alignment(vertical="center", wrap_text=ci in (8,9,10,11))
                        cell.font = Font(size=10)
                        if ci == 4 and sfill: cell.fill = sfill
                    ws.row_dimensions[row_idx].height = 16
                    row_idx += 1

            for i, w in enumerate([12,28,35,16,14,14,18,40,45,35,40], 1):
                ws.column_dimensions[get_column_letter(i)].width = w
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = f"A1:{get_column_letter(len(HEADERS))}1"
            wb.save(path)
            return {"ok": True, "path": path, "count": row_idx - 2}
        except Exception as e:
            import traceback
            return {"ok": False, "msg": str(e), "trace": traceback.format_exc()}

    # importar planilha no formato nativo (★ Importar Planilha na sidebar)
    def import_native_xlsx(self):
        try:
            path = open_file_dialog("Selecionar planilha",
                filetypes=[("Excel","*.xlsx *.xls"),("Todos","*.*")])
            if not path: return {"ok": False, "msg": "Cancelado"}
            try: import openpyxl
            except ImportError:
                return {"ok": False, "msg": "Instale openpyxl: pip install openpyxl"}

            wb = openpyxl.load_workbook(path, data_only=True)
            # Aceita qualquer aba: preferência para "Rastreamento", depois qualquer nome
            preferred = ["Rastreamento", "Planilha1"]
            ws = None
            for name in preferred:
                if name in wb.sheetnames: ws = wb[name]; break
            if ws is None: ws = wb.active

            # Detecta colunas pelo cabeçalho — funciona para qualquer formato
            cm = detect_col_map(ws)

            def gcol(row, field, default=""):
                """Pega valor pelo nome do campo usando o col_map."""
                idx = cm.get(field)
                if idx is None or idx >= len(row): return default
                v = row[idx]
                return default if v is None else v

            def gs(row, field, default=""):
                return safe_str(gcol(row, field, default))

            existing_orders = load_data()
            new_orders, updated_orders = [], []

            for row in ws.iter_rows(min_row=2, values_only=True):
                if not any(row): continue

                nf_raw  = gs(row, "nf")
                cliente = gs(row, "cliente")
                if not nf_raw and not cliente: continue

                nf_fmt = f"NF-{nf_raw}" if nf_raw and not str(nf_raw).upper().startswith("NF") else nf_raw

                # Rastreio: pode conter "codigo - chavenfE" — separar se necessário
                rastreio_raw = gs(row, "rastreio")
                chave_nfe    = ""
                # Se rastreio tem 44 dígitos = é chave NF-e, não código de rastreio
                if len(rastreio_raw) == 44 and rastreio_raw.isdigit():
                    chave_nfe = rastreio_raw
                    rastreio_raw = ""
                # Formato "NF - codigo": extrair código de rastreio real
                elif " - " in rastreio_raw:
                    parts = rastreio_raw.split(" - ", 1)
                    rastreio_raw = parts[1].strip() if len(parts) > 1 else rastreio_raw
                # Chave NF-e explícita na coluna dedicada
                chave_col = gs(row, "chavNfe")
                if not chave_nfe and len(chave_col) == 44 and chave_col.isdigit():
                    chave_nfe = chave_col

                # Construir observações combinando todos os campos de texto relevantes
                obs_parts = []
                for f in ("observacoes", "motivoAtraso", "obsExtra1", "obsExtra2"):
                    v = gs(row, f)
                    if v: obs_parts.append(v)
                obs_txt = " | ".join(obs_parts) if obs_parts else ""

                # Normalizar REENVIO: "NÃO" → "NAO"
                reenvio_val = gs(row, "reenvio", "NAO").upper()
                reenvio_val = reenvio_val.replace("Ã", "A").replace("Â", "A")  # NÃO → NAO
                if reenvio_val not in ("SIM", "S", "NAO"): reenvio_val = "NAO"

                # Status
                status_xlsx = gs(row, "statusXlsx").upper()
                sf_xlsx     = gs(row, "statusFinal").upper()
                if sf_xlsx in STATUS_FINAL_TO_SYS:
                    sys_status = STATUS_FINAL_TO_SYS[sf_xlsx]
                elif status_xlsx in STATUS_XLSX_TO_SYS:
                    sys_status = STATUS_XLSX_TO_SYS[status_xlsx]
                else:
                    sys_status = "em_transito"
                if reenvio_val in ("SIM", "S"): sys_status = "reenvio"

                # Transportadora normalizada
                trans_norm = norm_transportadora(gs(row, "transportadora"))

                # Plataforma: detectar pelo canal de venda ou prefixo do número do pedido
                canal     = gs(row, "plataforma")
                num_ped_raw = gcol(row, "numeroPedido")
                # Float inteiro (ex: 152247010410.0) → str sem .0
                if isinstance(num_ped_raw, float) and num_ped_raw == int(num_ped_raw):
                    num_ped = str(int(num_ped_raw))
                else:
                    num_ped = safe_str(num_ped_raw)
                # Adiciona prefixo LU- se for ID Magalu numérico (10-14 dígitos)
                if re.match(r"^\d{10,14}$", num_ped):
                    num_ped = "LU-" + num_ped
                # Detecta plataforma pelo formato do ID
                if num_ped.startswith("LU-"):
                    plataforma = "Magalu"
                elif num_ped.startswith(("701-","702-","703-","704-")) or re.match(r"^\d{3}-\d{7}-\d{7}$", num_ped):
                    plataforma = "Amazon"
                elif canal:
                    plataforma = canal
                else:
                    plataforma = "Outros"

                # Risco da planilha (pode não existir — padrão sem_risco)
                risco_raw = gs(row, "risco")
                risco_val = risco_raw if risco_raw in ("sem_risco","possivel_risco","emergencia") else "sem_risco"

                # ── Deduplicação ──────────────────────────────────────────
                existing = find_existing_order(existing_orders, nf=nf_fmt, chave_nfe=chave_nfe)
                if existing:
                    changed = False
                    def comp(field, val):
                        nonlocal changed
                        if val and not existing.get(field):
                            existing[field] = val; changed = True
                    comp("chavNfe",        chave_nfe)
                    comp("rastreio",       rastreio_raw)
                    comp("cidade",         gs(row, "cidade"))
                    comp("uf",             gs(row, "uf"))
                    comp("responsavel",    gs(row, "responsavel"))
                    comp("clienteAvisado", gs(row, "clienteAvisado"))
                    comp("rotaEntrega",    gs(row, "rotaEntrega"))
                    dr = gcol(row, "dataEntregaReal")
                    if dr: comp("dataEntregaReal", safe_date_iso(dr))
                    cf = gcol(row, "custoFrete")
                    if cf:
                        try: existing["custoFrete"] = existing.get("custoFrete") or float(str(cf).replace(",",".")); changed = True
                        except: pass
                    if changed:
                        updated_orders.append(existing["nf"])
                    continue

                # ── Novo pedido ───────────────────────────────────────────
                new_orders.append({
                    "id":              next_id(),
                    "nf":              nf_fmt,
                    "dataEnvio":       safe_date_iso(gcol(row, "dataEnvio")),
                    "reenvio":         reenvio_val,
                    "numeroPedido":    num_ped,
                    "cliente":         cliente,
                    "plataforma":      plataforma,
                    "transportadora":  trans_norm,
                    "previsaoPlataforma": safe_date_iso(gcol(row, "previsaoPlataforma")),
                    "previsaoEntrega":    safe_date_iso(gcol(row, "previsaoEntrega")),
                    "status":     sys_status,
                    "statusXlsx": status_xlsx,
                    "statusFinal": sf_xlsx,
                    "observacoes": [{"id":1,"texto":obs_txt,"data":"","usuario":"Importacao"}] if obs_txt else [],
                    "rotaEntrega":    gs(row, "rotaEntrega"),
                    "rastreio":       rastreio_raw,
                    "chavNfe":        chave_nfe,
                    "clienteAvisado": gs(row, "clienteAvisado"),
                    "dataContato":    (lambda v: safe_date_iso(v) if v else datetime.now().strftime("%d/%m/%Y %H:%M"))(gcol(row, "dataContato")),
                    "responsavel":    gs(row, "responsavel"),
                    "cidade":         gs(row, "cidade"),
                    "uf":             gs(row, "uf"),
                    "volumes": (lambda v: int(float(str(v).replace(",",".")) or 1) if v else 1)(gcol(row, "volumes")),
                    "peso": 0.0, "valor": 0.0,
                    "custoFrete":   safe_float(gcol(row, "custoFrete",   0)),
                    "custoReenvio": safe_float(gcol(row, "custoReenvio", 0)),
                    "custoExtra":   safe_float(gcol(row, "custoExtra",   0)),
                    "dataEntregaReal":  safe_date_iso(gcol(row, "dataEntregaReal")) if gcol(row, "dataEntregaReal") else "",
                    "diasEntregaReal":  (lambda v: int(v) if v and str(v).strip().lstrip("-").isdigit() else None)(gcol(row, "diasEntregaReal")),
                    "nfOriginal": "", "nfReenvioVinculada": "",
                    "risco": risco_val,
                    "chamados": [],
                    "historico": [{"data": datetime.now().strftime("%d/%m/%Y %H:%M"),
                                   "usuario": "Importacao", "acao": "Importado da planilha Excel"}],
                })
            wb.close()
            # Retornar todos os pedidos: existentes (já atualizados in-place) + novos
            all_orders = existing_orders + new_orders
            return {
                "ok": True, "count": len(new_orders),
                "updated": len(updated_orders),
                "orders": all_orders,
                "updatedNFs": updated_orders
            }
        except Exception as e:
            import traceback
            return {"ok": False, "msg": str(e), "trace": traceback.format_exc()}

    # exportar planilha no layout padrão (botão Salvar)
    def export_native_xlsx(self, orders):
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            return {"ok": False, "msg": "openpyxl nao instalado"}
        try:
            import tkinter as tk
            from tkinter import filedialog
            root = tk.Tk(); root.withdraw(); root.attributes("-topmost", True)
            path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel", "*.xlsx")],
                title="Salvar planilha",
                initialfile=f"LogSystem_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx"
            )
            root.destroy()
            if not path:
                return {"ok": False, "msg": "Cancelado"}

            # ── Helpers de estilo ────────────────────────────────────
            def fill(hex_):
                return PatternFill("solid", fgColor=hex_)
            def fnt(bold=False, color="1E293B", size=9):
                return Font(bold=bold, color=color, size=size, name="Arial")
            def aln(h="left", v="center", wrap=False):
                return Alignment(horizontal=h, vertical=v, wrap_text=wrap)
            thin = Side(style="thin", color="E2E8F0")
            def brd():
                return Border(left=thin, right=thin, top=thin, bottom=thin)

            # ── Paleta ───────────────────────────────────────────────
            HDR_BG="1E3A8A"; HDR_FG="FFFFFF"
            TTL_BG="0F172A"; TTL_FG="38BDF8"
            ROW0="FFFFFF";   ROW1="F8FAFC"
            GRN_B="DCFCE7";  GRN_F="166534"
            RED_B="FEE2E2";  RED_F="991B1B"
            AMB_B="FEF3C7";  AMB_F="92400E"
            BLU_B="DBEAFE";  BLU_F="1E40AF"
            PUR_B="EDE9FE";  PUR_F="5B21B6"
            GRY_B="F1F5F9";  GRY_F="475569"

            ST = {
                "entregue_prazo":  (GRN_B,GRN_F,"Entregue no Prazo"),
                "reenvio_prazo":   (GRN_B,GRN_F,"Reenvio Entregue"),
                "no_prazo":        (BLU_B,BLU_F,"No Prazo"),
                "em_transito":     (BLU_B,BLU_F,"Em Transito"),
                "em_rota":         (PUR_B,PUR_F,"Saiu para Entrega"),
                "reenvio":         (AMB_B,AMB_F,"Reenvio"),
                "risco_atraso":    (AMB_B,AMB_F,"Atencao: Prazo"),
                "em_atraso":       (RED_B,RED_F,"Em Atraso"),
                "entregue_atraso": (RED_B,RED_F,"Entregue Fora do Prazo"),
                "reenvio_atraso":  (RED_B,RED_F,"Reenvio Fora do Prazo"),
                "nao_entregue":    (RED_B,RED_F,"Nao Entregue"),
                "extraviado":      (GRY_B,GRY_F,"Extraviado"),
                "retorno":         (PUR_B,PUR_F,"Retorno"),
            }
            RC = {
                "emergencia":     (RED_B,RED_F,"EMERGENCIA"),
                "possivel_risco": (AMB_B,AMB_F,"POSSIVEL RISCO"),
                "sem_risco":      (GRN_B,GRN_F,"Sem Risco"),
            }

            def fmt_d(iso):
                if not iso: return ""
                s = str(iso)
                for fmt in ("%Y-%m-%dT%H:%M:%S","%Y-%m-%dT%H:%M","%Y-%m-%d"):
                    try: return datetime.strptime(s[:len(fmt)],fmt).strftime("%d/%m/%Y")
                    except: pass
                return s[:10] if "/" not in s else s[:10]

            def obs_txt(lst):
                if not lst or not isinstance(lst,list): return ""
                return " | ".join(o.get("texto","") for o in lst if o.get("texto"))

            today = datetime.now().date()
            gen   = datetime.now().strftime("%d/%m/%Y %H:%M")
            wb    = openpyxl.Workbook()

            # ════════════════════════════════════════════════════════
            # ABA 1 — PEDIDOS
            # ════════════════════════════════════════════════════════
            ws1 = wb.active; ws1.title = "Pedidos"
            ws1.sheet_view.showGridLines = False

            COLS = [
                ("DATA ENVIO",       13, "dataEnvio"),
                ("NOTA FISCAL",      13, "nf"),
                ("PEDIDO ID",        22, "numeroPedido"),
                ("CLIENTE",          28, "cliente"),
                ("CIDADE",           18, "cidade"),
                ("UF",                5, "uf"),
                ("PLATAFORMA",       13, "plataforma"),
                ("TRANSPORTADORA",   16, "transportadora"),
                ("RASTREIO",         22, "rastreio"),
                ("PRAZO PLAT.",      14, "previsaoPlataforma"),
                ("PRAZO TRANS.",     13, "previsaoEntrega"),
                ("ENTREGA REAL",     13, "dataEntregaReal"),
                ("DIAS",              6, "diasEntregaReal"),
                ("FRETE",            10, "custoFrete"),
                ("REENVIO$",         10, "custoReenvio"),
                ("EXTRA$",           10, "custoExtra"),
                ("VOLS",              5, "volumes"),
                ("RISCO",            15, "risco"),
                ("STATUS",           22, "status"),
                ("RESPONSAVEL",      15, "responsavel"),
                ("DATA CONTATO",     15, "dataContato"),
                ("OBS",              40, "observacoes"),
            ]
            NC = len(COLS)

            # Título
            ws1.merge_cells(f"A1:{get_column_letter(NC)}1")
            c = ws1["A1"]
            c.value = f"LogSystem Pro  |  {len(orders)} pedidos  |  Gerado em {gen}"
            c.fill=fill(TTL_BG); c.font=Font(bold=True,color=TTL_FG,size=11,name="Arial")
            c.alignment=aln("center"); ws1.row_dimensions[1].height=20

            # Cabeçalho
            for ci,(h,w,_) in enumerate(COLS,1):
                c = ws1.cell(row=2,column=ci,value=h)
                c.fill=fill(HDR_BG); c.font=fnt(True,HDR_FG,9)
                c.alignment=aln("center"); c.border=brd()
                ws1.column_dimensions[get_column_letter(ci)].width=w
            ws1.row_dimensions[2].height=18
            ws1.freeze_panes="A3"

            # Linhas
            for ri,o in enumerate(orders):
                rn = ri+3
                bg = ROW0 if ri%2==0 else ROW1
                st_key = o.get("status","")
                st_info = ST.get(st_key,(bg,"1E293B",st_key))
                rc_key  = o.get("risco","")
                rc_info = RC.get(rc_key,(bg,"1E293B",rc_key))

                for ci,(_,_,field) in enumerate(COLS,1):
                    raw = o.get(field,"")
                    if field=="observacoes":        val=obs_txt(raw)
                    elif field in ("dataEnvio","previsaoPlataforma","previsaoEntrega","dataEntregaReal"):
                        val=fmt_d(str(raw)) if raw else ""
                    elif field=="status":           val=st_info[2]
                    elif field=="risco":            val=rc_info[2]
                    elif field in ("custoFrete","custoReenvio","custoExtra"):
                        val=float(raw) if raw else 0.0
                    elif field=="diasEntregaReal":  val=int(raw) if raw else ""
                    elif field=="volumes":          val=int(raw) if raw else 1
                    else:                           val=str(raw) if raw else ""

                    c = ws1.cell(row=rn,column=ci,value=val)
                    c.border=brd(); c.font=fnt(color="1E293B")
                    c.alignment=aln("left",wrap=(field=="observacoes"))

                    if field=="status":
                        c.fill=fill(st_info[0]); c.font=fnt(True,st_info[1])
                        c.alignment=aln("center")
                    elif field=="risco":
                        c.fill=fill(rc_info[0]); c.font=fnt(True,rc_info[1])
                        c.alignment=aln("center")
                    elif field in ("previsaoPlataforma","previsaoEntrega"):
                        c.fill=fill(AMB_B if val else bg)
                        c.font=fnt(bool(val),AMB_F if val else "94A3B8")
                        c.alignment=aln("center")
                    elif field=="dataEntregaReal":
                        c.fill=fill(GRN_B if val else bg)
                        c.font=fnt(color=GRN_F if val else "94A3B8")
                        c.alignment=aln("center")
                    elif field in ("custoFrete","custoReenvio","custoExtra"):
                        c.number_format="R$ #,##0.00"; c.fill=fill(bg)
                        c.alignment=aln("right")
                    elif field=="diasEntregaReal" and isinstance(val,int) and val:
                        cc=GRN_B if val<=5 else AMB_B if val<=10 else RED_B
                        fc=GRN_F if val<=5 else AMB_F if val<=10 else RED_F
                        c.fill=fill(cc); c.font=fnt(True,fc)
                        c.alignment=aln("center")
                    elif field in ("nf","rastreio","numeroPedido"):
                        c.fill=fill(bg)
                        c.font=Font(size=9,color=BLU_F,name="Courier New")
                    else:
                        c.fill=fill(bg)

                ws1.row_dimensions[rn].height=15

            ws1.auto_filter.ref=f"A2:{get_column_letter(NC)}{len(orders)+2}"

            # ════════════════════════════════════════════════════════
            # ABA 2 — POR TRANSPORTADORA
            # ════════════════════════════════════════════════════════
            ws2 = wb.create_sheet("Por Transportadora")
            ws2.sheet_view.showGridLines = False

            from collections import defaultdict
            car = defaultdict(lambda:{"total":0,"ok":0,"fail":0,"custo":0.0,
                                      "atraso":0,"extrav":0,"retorno":0,"em_aberto":0})
            for o in orders:
                tr=o.get("transportadora","N/A") or "N/A"
                d=car[tr]; d["total"]+=1
                st=o.get("status","")
                d["custo"]+=float(o.get("custoFrete",0) or 0)
                if st in ("entregue_prazo","reenvio_prazo"): d["ok"]+=1
                elif st in ("entregue_atraso","reenvio_atraso","nao_entregue"): d["fail"]+=1
                elif st in ("em_atraso","risco_atraso"): d["atraso"]+=1
                if st=="extraviado": d["extrav"]+=1
                if st=="retorno":    d["retorno"]+=1
                dl=o.get("previsaoPlataforma") or o.get("previsaoEntrega","")
                if st in ("em_transito","no_prazo","risco_atraso","em_atraso","em_rota","reenvio"):
                    d["em_aberto"]+=1

            ws2.merge_cells("A1:I1")
            c=ws2["A1"]
            c.value=f"LogSystem Pro — Desempenho por Transportadora | {gen}"
            c.fill=fill(TTL_BG); c.font=Font(bold=True,color=TTL_FG,size=11,name="Arial")
            c.alignment=aln("center"); ws2.row_dimensions[1].height=20

            H2=[("TRANSPORTADORA",22),("TOTAL",9),("OK NO PRAZO",13),("FALHA",9),
                ("SLA %",9),("CUSTO FRETE",13),("EM ATRASO",11),("EXTRAVIADO",11),("EM ABERTO",11)]
            for ci,(h,w) in enumerate(H2,1):
                c=ws2.cell(row=2,column=ci,value=h)
                c.fill=fill(HDR_BG); c.font=fnt(True,HDR_FG,9)
                c.alignment=aln("center"); c.border=brd()
                ws2.column_dimensions[get_column_letter(ci)].width=w
            ws2.row_dimensions[2].height=18; ws2.freeze_panes="A3"

            for ri,(tr,d) in enumerate(sorted(car.items(),key=lambda x:-x[1]["total"])):
                rn=ri+3; bg=ROW0 if ri%2==0 else ROW1
                ok=d["ok"]; fail=d["fail"]; tot=ok+fail
                sla=round(ok/tot*100) if tot else 0
                sc_b=GRN_B if sla>=90 else AMB_B if sla>=70 else RED_B
                sc_f=GRN_F if sla>=90 else AMB_F if sla>=70 else RED_F
                vals=[tr,d["total"],ok,fail,f"{sla}%",d["custo"],d["atraso"],d["extrav"],d["em_aberto"]]
                for ci,val in enumerate(vals,1):
                    c=ws2.cell(row=rn,column=ci,value=val)
                    c.border=brd(); c.fill=fill(bg); c.font=fnt(ci==1,color="1E293B")
                    c.alignment=aln("left" if ci==1 else "center")
                    if ci==5:  c.fill=fill(sc_b); c.font=fnt(True,sc_f)
                    elif ci==6: c.number_format="R$ #,##0.00"; c.alignment=aln("right")
                    elif ci in(4,7,8) and isinstance(val,int) and val>0:
                        c.fill=fill(RED_B); c.font=fnt(True,RED_F)
                ws2.row_dimensions[rn].height=15

            # ════════════════════════════════════════════════════════
            # ABA 3 — ATRASOS E RISCOS (para reunião com transportadora)
            # ════════════════════════════════════════════════════════
            ws3 = wb.create_sheet("Atrasos e Riscos")
            ws3.sheet_view.showGridLines = False

            risk_orders=[o for o in orders if
                o.get("risco") in ("emergencia","possivel_risco") or
                o.get("status") in ("em_atraso","nao_entregue","risco_atraso","extraviado","retorno")]

            ws3.merge_cells("A1:K1")
            c=ws3["A1"]
            c.value=f"LogSystem Pro — Atrasos e Riscos | {len(risk_orders)} registros | {gen}"
            c.fill=fill("7F1D1D"); c.font=Font(bold=True,color="FECACA",size=11,name="Arial")
            c.alignment=aln("center"); ws3.row_dimensions[1].height=20

            H3=[("NF",13),("PEDIDO ID",22),("CLIENTE",28),("PLATAFORMA",13),
                ("TRANSPORTADORA",16),("PRAZO PLAT.",14),("PRAZO TRANS.",13),
                ("DIAS ATRASO",11),("RISCO",15),("STATUS",22),("RASTREIO",22)]
            for ci,(h,w) in enumerate(H3,1):
                c=ws3.cell(row=2,column=ci,value=h)
                c.fill=fill("991B1B"); c.font=fnt(True,"FFFFFF",9)
                c.alignment=aln("center"); c.border=brd()
                ws3.column_dimensions[get_column_letter(ci)].width=w
            ws3.row_dimensions[2].height=18; ws3.freeze_panes="A3"

            for ri,o in enumerate(risk_orders):
                rn=ri+3; bg=ROW0 if ri%2==0 else ROW1
                st_key=o.get("status",""); rc_key=o.get("risco","")
                st_i=ST.get(st_key,(bg,"1E293B",st_key))
                rc_i=RC.get(rc_key,(bg,"1E293B",rc_key))
                dl=o.get("previsaoPlataforma") or o.get("previsaoEntrega","")
                atr=""
                if dl:
                    try:
                        dl_d=datetime.strptime(dl[:10],"%Y-%m-%d").date()
                        diff=(today-dl_d).days
                        atr=f"+{diff}d" if diff>0 else ("Hoje" if diff==0 else f"{diff}d")
                    except: pass

                vals=[o.get("nf",""),o.get("numeroPedido",""),o.get("cliente",""),
                      o.get("plataforma",""),o.get("transportadora",""),
                      fmt_d(dl),fmt_d(o.get("previsaoEntrega","")),
                      atr,rc_i[2],st_i[2],o.get("rastreio","")]
                for ci,val in enumerate(vals,1):
                    c=ws3.cell(row=rn,column=ci,value=val)
                    c.border=brd(); c.font=fnt(color="1E293B")
                    c.alignment=aln("left" if ci<=3 else "center")
                    if ci==9:   c.fill=fill(rc_i[0]); c.font=fnt(True,rc_i[1])
                    elif ci==10: c.fill=fill(st_i[0]); c.font=fnt(True,st_i[1])
                    elif ci==8 and val and str(val).startswith("+"):
                        c.fill=fill(RED_B); c.font=fnt(True,RED_F)
                    elif ci in (1,2): c.fill=fill(bg); c.font=Font(size=9,color=BLU_F,name="Courier New")
                    else: c.fill=fill(bg)
                ws3.row_dimensions[rn].height=15

            ws3.auto_filter.ref=f"A2:K{len(risk_orders)+2}"

            wb.save(path)
            return {"ok": True, "count": len(orders), "path": path}

        except Exception as e:
            import traceback
            return {"ok": False, "msg": str(e), "trace": traceback.format_exc()}

    # exportar abas da tela Analíticos
    def export_analytics_xlsx(self, data):
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            return {"ok": False, "msg": "openpyxl nao instalado"}
        try:
            fname = f"analitico_logistica_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            path = open_file_dialog("Salvar relatorio analitico",
                filetypes=[("Excel","*.xlsx")], save=True, save_name=fname)
            if not path: return {"ok": False, "msg": "Cancelado"}
            wb = openpyxl.Workbook()
            hfill = PatternFill("solid", fgColor="1F3864")
            hfont = Font(bold=True, color="FFFFFF", size=10)
            thin = Side(border_style="thin", color="CCCCCC")
            bdr = Border(left=thin, right=thin, top=thin, bottom=thin)
            def make_sheet(wb, title, headers, rows_data):
                ws = wb.create_sheet(title=title)
                for c, h in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=c, value=h)
                    cell.fill=hfill; cell.font=hfont
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border=bdr
                for ri, row in enumerate(rows_data, 2):
                    for ci, val in enumerate(row, 1):
                        cell = ws.cell(row=ri, column=ci, value=val)
                        cell.border=bdr; cell.alignment=Alignment(vertical="center")
                ws.freeze_panes="A2"
                return ws
            if "byUF" in data:
                make_sheet(wb,"Por Estado",
                    ["Estado","Total Pedidos","Entregues","No Prazo","Com Atraso",
                     "Media Dias Entrega","Melhor Transportadora","% Atraso","Custo Medio Frete"],
                    data["byUF"])
            if "byCarrier" in data:
                make_sheet(wb,"Por Transportadora",
                    ["Transportadora","Total","Entregues","No Prazo","Atrasos",
                     "SLA %","Media Dias","Custo Total","Custo Medio"],
                    data["byCarrier"])
            if "cost" in data:
                make_sheet(wb,"Custo x Efetivo",
                    ["Transportadora","Custo Frete Total","Custo Reenvio","Custo Extra",
                     "Custo Efetivo Total","Diferenca","% Overhead"],
                    data["cost"])
            if "reenvios" in data:
                make_sheet(wb,"Reenvios",
                    ["NF Original","NF Reenvio","Cliente","Transportadora",
                     "Custo Original","Custo Reenvio","Custo Extra","Status"],
                    data["reenvios"])
            if "Sheet" in wb.sheetnames: del wb["Sheet"]
            wb.save(path)
            return {"ok": True, "path": path}
        except Exception as e:
            return {"ok": False, "msg": str(e)}

    # importar CSV genérico com mapeamento de colunas
    def import_file(self):
        try:
            path = open_file_dialog("Selecionar arquivo",
                filetypes=[("Planilhas","*.csv *.xlsx *.xls"),("CSV","*.csv"),("Excel","*.xlsx *.xls"),("Todos","*.*")])
            if not path: return {"ok": False, "msg": "Cancelado"}
            ext = os.path.splitext(path)[1].lower()
            if ext == ".csv":
                rows = []
                with open(path,"r",encoding="utf-8-sig",errors="replace") as f:
                    sample = f.read(2048); f.seek(0)
                    try: dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
                    except: dialect = csv.excel
                    for row in csv.reader(f, dialect): rows.append(row)
                if not rows: return {"ok": False, "msg": "Arquivo vazio"}
                return {"ok": True, "headers": rows[0], "rows": rows[1:], "format": "csv"}
            elif ext in (".xlsx",".xls"):
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(path, data_only=True); ws = wb.active
                    rows = [[str(c) if c is not None else "" for c in row] for row in ws.iter_rows(values_only=True)]
                    wb.close()
                    if not rows: return {"ok": False, "msg": "Arquivo vazio"}
                    return {"ok": True, "headers": rows[0], "rows": rows[1:], "format": "xlsx"}
                except ImportError:
                    return {"ok": False, "msg": "Instale openpyxl"}
            return {"ok": False, "msg": "Formato nao suportado"}
        except Exception as e:
            return {"ok": False, "msg": str(e)}

    # exportar CSV simplificado
    def export_csv(self, orders):
        try:
            fname = f"rastreamento_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
            path = open_file_dialog("Salvar CSV", filetypes=[("CSV","*.csv"),("Todos","*.*")],
                save=True, save_name=fname)
            if not path: return {"ok": False, "msg": "Cancelado"}
            output = io.StringIO()
            w = csv.writer(output, delimiter=";")
            w.writerow(["DATA DO PEDIDO","NOTA FISCAL","CHAVE NF-e","REENVIO?","N DO PEDIDO","CLIENTE",
                        "CANAL DE VENDA","TRANSPORTADORA","PRAZO PLATAFORMA","PRAZO TRANSPORTADORA",
                        "STATUS DO PEDIDO","OBSERVACOES","ROTA?","RASTREIO","CLIENTE AVISADO",
                        "DATA CONTATO","RESPONSAVEL","STATUS FINAL","CIDADE","ESTADO",
                        "DATA ENTREGA REAL","DIAS ENTREGA","CUSTO FRETE","CUSTO REENVIO","CUSTO EXTRA","RISCO"])
            for o in orders:
                obs = " | ".join(x.get("texto","") for x in (o.get("observacoes") or []) if x.get("texto"))
                st, sf = STATUS_SYS_TO_XLSX.get(o.get("status","em_transito"),("NO PRAZO","EM VIAGEM"))
                fd = lambda k: (o.get(k,"") or "")[:10]
                w.writerow([fd("dataEnvio"), str(o.get("nf","")).replace("NF-",""),
                    o.get("chavNfe",""), o.get("reenvio","NAO"), o.get("numeroPedido",""), o.get("cliente",""),
                    o.get("plataforma",""), o.get("transportadora",""),
                    fd("previsaoPlataforma"), fd("previsaoEntrega"),
                    st, obs, o.get("rotaEntrega",""), o.get("rastreio",""),
                    o.get("clienteAvisado",""), fd("dataContato"), o.get("responsavel",""), sf,
                    o.get("cidade",""), o.get("uf",""),
                    fd("dataEntregaReal"), o.get("diasEntregaReal",""),
                    o.get("custoFrete",0), o.get("custoReenvio",0), o.get("custoExtra",0),
                    o.get("risco","sem_risco")])
            with open(path,"w",encoding="utf-8-sig",newline="") as f:
                f.write(output.getvalue())
            return {"ok": True, "path": path}
        except Exception as e:
            return {"ok": False, "msg": str(e)}


# Acha colunas da planilha pelo cabeçalho — funciona com layout variado
def detect_col_map(ws):
    """Map field names to 0-indexed column positions by reading the header row.
    Works for any spreadsheet format regardless of column order."""
    import unicodedata
    def norm(s):
        if s is None: return ""
        s = str(s).strip().upper()
        s = unicodedata.normalize("NFD", s)
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        # remove special chars except alphanumerics and spaces
        s = "".join(ch if ch.isalnum() or ch == " " else " " for ch in s)
        return " ".join(s.split())  # collapse whitespace
    # Each field: list of normalized aliases (first match wins)
    FIELDS = {
        "nf":               ["NOTA FISCAL", "NF"],
        "dataEnvio":        ["DATA DO PEDIDO", "DATA PEDIDO", "DATA"],
        "reenvio":          ["REENVIO"],
        "numeroPedido":     ["N DO PEDIDO", "NUMERO DO PEDIDO"],
        "cliente":          ["CLIENTE"],
        "plataforma":       ["CANAL DE VENDA", "CANAL VENDA", "PLATAFORMA", "CANAL"],
        "transportadora":   ["TRANSPORTADORA", "TRANS"],
        "previsaoPlataforma":["PRAZO PLATAFORMA", "PRAZO ENTREGA PLATAFORMA"],
        "previsaoEntrega":  ["PRAZO TRANSPORTADORA", "PRAZO ENTREGA TRANSPORTADORA"],
        "statusXlsx":       ["STATUS DO PEDIDO", "STATUS PEDIDO"],
        "motivoAtraso":     ["MOTIVO DO ATRASO", "MOTIVO ATRASO"],
        "obsExtra1":        ["DATA CONTATO TRANSPORTADORA", "CONTATO TRANSPORTADORA"],
        "obsExtra2":        ["RETORNO DA TRANSPORTADORA", "RETORNO TRANSPORTADORA"],
        "observacoes":      ["OBSERVACOES", "OBSERVA"],  # partial match via 'in'
        "rotaEntrega":      ["ROTA DE ENTREGA", "ROTA ENTREGA"],
        "rastreio":         ["RASTREIO"],
        "clienteAvisado":   ["CLIENTE AVISADO"],
        "dataContato":      ["DATA CONTATO CLIENTE"],  # sem "DATA CONTATO" genérico para não colidir com TRANSPORTADORA
        "responsavel":      ["RESPONSAVEL PELO ATENDIMENTO", "RESPONSAVEL"],
        "statusFinal":      ["STATUS FINAL"],
        "chavNfe":          ["CHAVE NF E", "CHAVE NFE"],
        "cidade":           ["CIDADE"],
        "uf":               ["ESTADO", "UF"],
        "dataEntregaReal":  ["DATA ENTREGA REAL"],
        "diasEntregaReal":  ["DIAS ENTREGA REAL"],
        "custoFrete":       ["CUSTO FRETE"],
        "custoReenvio":     ["CUSTO REENVIO"],
        "custoExtra":       ["CUSTO EXTRA"],
        "volumes":          ["VOLUMES"],
        "risco":            ["RISCO"],
        "chamados":         ["CHAMADOS"],
    }
    col_map = {}
    for col in range(1, ws.max_column + 1):
        raw = ws.cell(1, col).value
        if not raw: continue
        hn = norm(raw)
        for field, aliases in FIELDS.items():
            if field in col_map: continue  # already mapped
            for alias in aliases:
                an = norm(alias)
                # Exact match ou header começa com o alias (prefixo)
                # Nunca usar "an in hn" (muito permissivo) ou "an.endswith(hn)" (errado)
                if hn == an or hn.startswith(an):
                    col_map[field] = col - 1  # 0-indexed
                    break
    return col_map

def norm_transportadora(raw):
    """Padroniza nome da transportadora (CORREIOS → Correios, etc.)."""
    if not raw: return ""
    r = str(raw).strip().upper()
    mapping = {
        "CORREIOS": "Correios", "ECT": "Correios",
        "J&T EXPRESS": "J&T Express", "J&T": "J&T Express", "JNT": "J&T Express",
        "BRASPRESS": "Braspress",
        "SAO MIGUEL": "Sao Miguel", "SÃO MIGUEL": "Sao Miguel",
        "AZUL CARGO": "Azul Cargo", "AZUL": "Azul Cargo",
        "JADLOG": "Jadlog", "LOGGI": "Loggi", "SEQUOIA": "Sequoia",
        "TOTAL EXPRESS": "Total Express", "MAGALU ENTREGAS": "Magalu Entregas",
        "SHOPEE EXPRESS": "Shopee Express", "DIRECT": "Direct",
    }
    for key, val in mapping.items():
        if r == key or r.startswith(key): return val
    # Title-case fallback
    return raw.strip().title()

def _startup_checks():
    """Deps, banco SQLite e app.html — aborta com mensagem se algo falhar."""
    database.configure(BASE_DIR)
    auth.load_users(USERS_FILE)

    ok, msg = bootstrap.ensure_dependencies(auto_install=True)
    if not ok:
        print(f"[ERRO] {msg}")
        sys.exit(1)
    print(f"[LogSystem] {msg}")

    ok, msg = database.verificar_conexao()
    if not ok:
        print(f"[ERRO DB] {msg}")
        sys.exit(1)
    print(f"[LogSystem] {msg}")

    try:
        migrated = database.inicializar_banco()
        if migrated:
            print(f"[LogSystem] Migrados {migrated} pedido(s) de data.json para SQLite")
    except Exception as exc:
        print(f"[ERRO DB] Falha ao inicializar banco: {exc}")
        sys.exit(1)

    has_backend, backend_msg = bootstrap.check_optional_backend()
    if not has_backend:
        print(f"[AVISO] {backend_msg}")

    if not os.path.exists(HTML_FILE):
        print(f"[ERRO] Interface não encontrada: {HTML_FILE}")
        sys.exit(1)


if __name__ == "__main__":
    _startup_checks()

    import webview

    # backup a cada 5 min enquanto o app estiver aberto
    bt = threading.Thread(target=backup_loop, daemon=True)
    bt.start()
    print(f"[LogSystem v{APP_VERSION}] Backup ativo → {BACKUP_DIR}")

    api = API()
    window = webview.create_window(
        title=f"LogSystem Pro v{APP_VERSION}",
        url=f"file://{HTML_FILE}", js_api=api,
        width=1500, height=920, min_size=(1200, 720),
        background_color="#000000", text_select=False)
    webview.start(debug=False)
